#!/usr/bin/env python3
# Generates DRC_Revision_Package.docx for Mohit's Garhwal Malta thesis revisions.
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- base styles ----------
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def H(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
    return p

def para(text, italic=False, bold=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    return p

def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)
    return p

def source(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    p.paragraph_format.space_after = Pt(10)
    return p

def make_table(headers, rows, widths=None, header_bg='1F3B57', font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        set_cell_bg(hdr[i], header_bg)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
        if ri % 2 == 1:
            for c in cells:
                set_cell_bg(c, 'EEF2F6')
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

# ============================================================= TITLE
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Doctoral Research Committee (DRC) Revision Package')
r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(0x1F,0x3B,0x57)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Optimisation of Malta (Citrus sinensis) Distribution Channels in the Garhwal Region of Uttarakhand')
r.italic = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Researcher: Mohit Gundwal   |   Discipline: Management   |   Prepared in response to DRC observations')
r.font.size = Pt(10); r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0x55,0x55,0x55)

doc.add_paragraph()

# ---- integrity / usage note ----
note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(6)
rr = note.add_run('How to use this document. ')
rr.bold = True; rr.font.size = Pt(10); rr.font.name='Times New Roman'
rr2 = note.add_run(
 'The text below is drafting scaffolding prepared from your own thesis, expert-feedback data, and geo-tagged field records. '
 'Read it as a working draft, not a final submission. Before you paste any part into the thesis or the DRC response, do three things: '
 '(1) verify every number, coordinate, name and affiliation against your primary records; '
 '(2) replace each item marked [CONFIRM] or [INSERT] with your own verified detail; and '
 '(3) revise the wording into your own voice so it reads as your work and matches the rest of the thesis. '
 'Passages are written in plain scholarly English on purpose so they blend with your existing chapters after your own editing pass.')
rr2.font.size = Pt(10); rr2.font.name='Times New Roman'

doc.add_page_break()

# ============================================================= SECTION 1
H('Section 1 — Region-Specific Outcomes and Conclusions (Garhwal Region)', level=1)
para('Suggested placement: Chapter 8, as a new sub-section 8.1a "Region-Specific Conclusions for the Garhwal Malta Belt", '
     'immediately after Table 8.1. It converts the study\'s general conclusions into statements tied explicitly to the four study districts.',
     italic=True, size=10)

H('8.1a Region-specific conclusions for the Garhwal Malta belt', level=2)
para('The conclusions of this study are not offered as generic supply-chain propositions; they are specific to the sweet-orange '
     '(Malta, Citrus sinensis) economy of the Garhwal Himalaya and to the four districts in which the field evidence was gathered '
     '\u2014 Pauri Garhwal, Tehri Garhwal, Rudraprayag and Chamoli. The following region-specific conclusions follow directly from '
     'the integrated Fuzzy AHP\u2013Fuzzy ELECTRE I analysis (Chapters 5 and 6), the stage-criticality outcomes, and the '
     'transportation optimisation model (Section 6.11), read against the terrain and market geography of the Garhwal belt.',
     align='justify')

H('Conclusion 1: In the Garhwal Malta channel, cost is the binding constraint on grower return', level=3)
para('For the Garhwal belt specifically, the criterion weights derived by Fuzzy AHP place cost decisively first '
     '(0.519), ahead of quality (0.308) and time (0.173). In a low-value, single-season, perishable hill crop moved over fragile '
     'mountain roads, it is the accumulation of many small costs \u2014 harvesting, grading, packing, head-load and vehicle haulage '
     'from dispersed orchards \u2014 that most visibly erodes the grower\u2019s share of the consumer rupee. The regional implication '
     'is that price-side measures alone (such as the C-grade Malta MSP of \u20b910/kg in 2024\u201325) will under-perform unless '
     'they are paired with cost-side relief at and just beyond the orchard gate.', align='justify')

H('Conclusion 2: Criticality in the Garhwal channel is concentrated upstream, at the Farm and the Mandi', level=3)
para('The Fuzzy ELECTRE I outranking for the Garhwal channel is Farm \u227b Mandi \u227b Transport \u227b Village Trader \u227b Retail. '
     'The Farm stage is the undominated kernel (Relative Criticality Index = 100) and the Mandi second (RCI = 62); together they '
     'account for roughly four-fifths of total channel criticality. For a region as dispersed and infrastructure-poor as Garhwal, '
     'this concentration is the single most useful planning result: scarce public and institutional effort should be directed first '
     'to the orchard interface and the wholesale-aggregation node, where value is both created and most easily lost.', align='justify')

H('Conclusion 3: The ranking is a structural feature of the Garhwal belt, not an artefact of the weights', level=3)
para('The stage ranking held invariant under \u00b115 and \u00b130 per cent perturbation of every criterion weight, and the pairwise '
     'comparison returned a consistency ratio of 0.0079. For the region this means the Farm-and-Mandi priority is a genuine '
     'structural characteristic of the Garhwal Malta channel and can be acted upon with confidence, rather than a fragile outcome '
     'that shifts with analytical assumptions.', align='justify')

H('Conclusion 4: A region-specific, least-cost dispatch geometry exists for Garhwal Malta', level=3)
para('The balanced transportation model over the four producing districts and the principal terminal markets, solved by Vogel\u2019s '
     'Approximation Method with a MODI optimality check, yields a least-cost dispatch plan whose total freight cost of '
     '\u20b934.71 million is about 7.3 per cent below a naive allocation. The regional reading is that each Garhwal district has a '
     'market to which it is most competitively matched, and that consolidating dispersed orchard output at decentralised '
     'collection-and-grading points before optimised, hub-and-spoke dispatch is the cost-minimising geometry for the belt.',
     align='justify')

H('Conclusion 5: District-level texture within the Garhwal frame', level=3)
para('While the channel-level ranking holds across the belt, the four districts differ in terrain, connectivity and Malta maturity, '
     'and the field evidence supports the following district-specific readings. [CONFIRM each against your district-level tabulations.]',
     align='justify')
make_table(
    ['District (sample n)', 'Terrain / access', 'Malta status', 'Region-specific conclusion for the district'],
    [
     ['Pauri Garhwal (311)', 'Mid-altitude, moderate connectivity', 'Established growing pockets',
      'Relatively better road access makes Pauri the most suitable district for early FPO-led aggregation and direct-to-retail linkage; farm-gate grading support yields quick cost relief.'],
     ['Tehri Garhwal (333)', 'Steep valleys, dispersed orchards', 'Significant production',
      'Orchard dispersion and steep valleys make near-farm collection-and-grading centres and reusable ventilated crates the highest-return interventions; head-load and first-mile cost dominate.'],
     ['Rudraprayag (316)', 'High-altitude, constrained roads', 'Emerging production',
      'As an emerging belt, Rudraprayag benefits most from feeder-road linkage and panchayat/village-level aggregation to prevent distress sales before the channel matures.'],
     ['Chamoli (340)', 'Remote, long market distances', 'Substantial production',
      'Longest market distances and landslide-prone routes make transport risk and cold-chain/reefer routing more binding here than elsewhere; decentralised processing (juice/value addition) is most justified.'],
    ],
    widths=[1.4, 1.5, 1.2, 3.2])
source('Source: Author, derived from Tables 4.2 and 4.5, the Fuzzy ELECTRE I ranking (Chapter 6), and district field enumeration. '
       'District terrain/Malta descriptors reproduced from Table 4.2.')

H('Conclusion 6: The regional opportunity is to convert upstream loss into retained hill-farmer income', level=3)
para('Read positively, the same upstream stages that leak value can, if strengthened, turn Malta from an under-exploited hill fruit '
     'into a flagship of the Garhwal horticulture economy. Because grading, packing and processing are activities in which rural '
     'women participate heavily \u2014 as the field observation of SHG-led Malta processing at Jantadevi (Chaubatta Khal, Pauri Garhwal) '
     'and the SHG enterprise sessions at Jakhnidhar (Tehri Garhwal) confirm \u2014 retaining value upstream also widens women\u2019s '
     'economic participation and helps slow distress out-migration from the hills.', align='justify')

# ============================================================= SECTION 2
doc.add_page_break()
H('Section 2 — Presentation of Geo-Tagged Field Photographs (Academic Format)', level=1)
para('Suggested placement: a short methodological note in Chapter 4 (Section 4.6.1, Primary data), with the full set of images placed '
     'in an Annexure titled "Annexure C: Geo-Tagged Field Photographs (Photographic Evidence of Field Work)". Representative plates '
     'may be embedded in Chapter 4 or 5 where they illustrate a specific point.', italic=True, size=10)

H('2.1 Governing conventions', level=2)
para('Geo-tagged photographs are treated as figures and follow the thesis\u2019s existing figure conventions, with four additions '
     'required for spatial evidence: (i) a location name, (ii) decimal-degree coordinates, (iii) the date and time of capture, and '
     '(iv) the datum/source of the geotag. Recommended rules:', align='justify')
numbered('Numbering. Use a dedicated "Plate" series for photographs (Plate 4.1, Plate 4.2 \u2026) so they are distinct from analytical '
         'figures and charts. If your institute requires a single figure series, number them continuously as Figures within the annexure '
         '(e.g., Figure C.1, Figure C.2).')
numbered('Caption position. Place the caption immediately below the photograph, left-aligned, in a smaller font than body text '
         '(typically 9\u201310 pt), consistent with all other figures in the thesis.')
numbered('Caption content, in fixed order: Plate number \u2192 short descriptive title \u2192 location (village, block, district) \u2192 '
         'coordinates in decimal degrees \u2192 date and time of capture \u2192 activity/relevance.')
numbered('Coordinates. Report latitude and longitude in decimal degrees to six places, with the datum (WGS 84, as used by GPS/GPS-camera '
         'applications). Optionally add altitude if recorded.')
numbered('Attribution and ethics. Add "Source: Author\u2019s field survey, [year]." Where individuals are identifiable, state in the '
         'methodology that verbal informed consent for photography was obtained (link to Section 4.15, Ethical Considerations). '
         'Faces of respondents may be blurred if your ethics protocol requires anonymity.')
numbered('Cross-referencing. Refer to each plate in the text at least once (e.g., "field observation at Pokhari, Chamoli, confirmed '
         'landslide-induced road blockage during transit (Plate C.7)").')
numbered('Master list. Provide a "List of Plates" in the front matter and a summary geo-tag table at the head of the annexure '
         '(see Table C.1 template below), so examiners can locate every site spatially.')

H('2.2 Recommended caption template', level=2)
p = doc.add_paragraph()
r = p.add_run('Plate 4.1  Administration of the structured questionnaire to a village retailer, Sabli (Chamba block), '
              'Tehri Garhwal. Coordinates: 30.320729\u00b0 N, 78.407015\u00b0 E (WGS 84). Captured 15 June 2025, 11:21 IST. '
              'The interaction illustrates primary data collection at the retail node of the Malta channel.')
r.italic = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
p2 = doc.add_paragraph()
r = p2.add_run('Source: Author\u2019s field survey, 2025.')
r.italic = True; r.font.size = Pt(8); r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0x55,0x55,0x55)

H('2.3 Table C.1 (template): Register of geo-tagged field photographs', level=2)
para('Populate this from your image metadata. The rows below are pre-filled from the photographs you supplied; verify each entry, '
     'confirm block names marked [CONFIRM], and add any images not listed.', size=10, italic=True)
make_table(
    ['Plate', 'Location (village, block, district)', 'Latitude (\u00b0N)', 'Longitude (\u00b0E)', 'Date & time (IST)', 'Field activity / relevance'],
    [
     ['C.1', 'Sabli / Ranichauri, Chamba [CONFIRM] block, Tehri Garhwal', '30.320729', '78.407015', '15-06-2025, 11:21', 'Structured questionnaire with retailer (retail node)'],
     ['C.2', 'Sabli / Ranichauri, Tehri Garhwal', '30.320729', '78.407015', '15-06-2025, 11:31', 'Questionnaire administration and instrument review'],
     ['C.3', 'Sabli / Ranichauri, Tehri Garhwal', '30.320729', '78.407015', '15-06-2025, 11:41', 'Field observation of orchard-to-road terrain and access'],
     ['C.4', 'Tharsalgaon, Sabli, Tehri Garhwal', '30.333418', '78.404446', '16-06-2025, 15:33', 'Interview with an elderly woman Malta grower (farm node)'],
     ['C.5', 'Jakhnidhar block, Tehri Garhwal', '30.334093', '78.510676', '26-12-2025 [CONFIRM]', 'Women SHG enterprise-development session (RSETI/NSSBDS)'],
     ['C.6', 'Petouri, Pratapnagar block, Tehri Garhwal', '30.482783', '78.476028', '22-06-2025, 11:17', 'Interview on landslide-prone feeder road (transport hazard)'],
     ['C.7', 'Kainyur / Raull, Thailisain block, Pauri Garhwal', '30.018682', '79.059656', '16-06-2025, 10:06', 'Researcher at field site; grower household approach'],
     ['C.8', 'Kainyur / Raull, Thailisain block, Pauri Garhwal', '30.018682', '79.059656', '16-06-2025, 11:48', 'Group interaction with a grower family (farm node)'],
     ['C.9', 'Jantadevi, Chaubatta Khal block, Pauri Garhwal', '29.943948', '78.823514', '05-10-2025, 17:26', 'SHG-led Malta peeling/processing (value addition)'],
     ['C.10', 'Jantadevi, Chaubatta Khal block, Pauri Garhwal', '29.943948', '78.823514', '05-10-2025, 14:17', 'SHG group discussion and product handling'],
     ['C.11', 'Sankri / Nail, Pokhari block, Chamoli', '30.382558', '79.212606', '28-06-2025, 12:15', 'Questionnaire with village traders (village-trader node)'],
     ['C.12', 'Sankri / Nail (SH-34), Pokhari block, Chamoli', '30.382558', '79.212606', '10-11-2025 [CONFIRM]', 'Interview on landslide-affected State Highway (transport risk)'],
     ['C.13', 'Rupung, Molta, Joshimath block, Chamoli', '30.482226', '79.519730', '22-06-2025 [CONFIRM time]', 'Interview with a woman grower in a remote orchard settlement'],
    ],
    widths=[0.5, 2.3, 0.8, 0.85, 1.05, 2.1], font_size=8)
source('Source: Author\u2019s field survey, 2025. Coordinates and timestamps read from GPS-camera metadata; datum WGS 84.')

H('2.4 Two data-integrity checks before submission', level=2)
bullet('Duplicate/geotag mismatch: two visually similar "man-with-umbrella on a cut-slope road" images carry different geotags '
       '(Sankri, Chamoli, 30.382558\u00b0/79.212606\u00b0 and Petouri, Tehri, 30.482783\u00b0/78.476028\u00b0) and different dates. '
       'Confirm each was taken at the stated site, or the committee may query it. Keep only correctly tagged originals.')
bullet('Date sanity: a few timestamps fall in the winter/late-2025 window (e.g., 26-12-2025, 05-10-2025, 10-11-2025). Confirm these '
       'match your fieldwork calendar; the Malta harvest season (roughly November\u2013January) is consistent with winter processing images, '
       'but the committee will expect the dates to align with your stated data-collection period.')

# ============================================================= SECTION 3
doc.add_page_break()
H('Section 3 — Expert Validation of Factors (Content-Validity Panel)', level=1)
para('Suggested placement: Chapter 4, Section 4.7.3 (Content validity and expert review), with the full table moved to '
     '"Annexure B: Expert Validation Panel and Factor-Level Feedback". The factor codes follow your five-stage \u00d7 three-dimension '
     'scheme: F = Farm, M = Mandi, T = Village Trader, TR = Transport, R = Retail; COST / TIME / QUAL denote the dimension.',
     italic=True, size=10)

H('4.7.3 Expert validation panel', level=2)
para('Content validity of the survey instrument and of the initial factor pool was established through a purposive panel of 45 domain '
     'experts drawn from horticulture science and academia, the state horticulture and agriculture departments, APMC/mandi boards, '
     'FPOs and cooperatives, Eco-Development Committees (EDCs), transport and logistics associations, and women SHG federations across '
     'the Garhwal districts. Each expert was asked three questions: (Q1) additional factors to be added; (Q2) factors judged '
     'irrelevant, redundant, or better merged; and (Q3) general comments on the Malta channel. Their feedback drove the reduction of '
     'the initial pool of 144 candidate factors to the final 82 retained factors (Section 4.11) and the merging or elimination of the '
     'items listed under "Factor(s) validated / flagged" below.', align='justify')

experts = [
 ['E01','Executive Member, EDC Valley of Flowers, Chamoli','Community-based eco-tourism & farm-gate supply','Transport hazard factor (added); R_COST_5 (flagged low-relevance)','Add explicit hill-terrain hazard risk (landslides, road blockages at peak harvest) under Transport; retail rent less relevant for orchard-gate direct sale; EDC-run collection points can cut transit spoilage ~25%.'],
 ['E02','Member, EDC Valley of Flowers, Chamoli','First-mile / head-load logistics & packaging','F_COST_2 + F_COST_4 (merge)','Add mule/manual head-load cost from non-roadhead orchards to vehicle points; merge crate and packing-material cost into one packaging-overhead variable; ventilated crates needed for high-hill to plains haul.'],
 ['E03','Member, EDC Valley of Flowers, Chamoli','Farm-gate pricing & payment terms','M_TIME_3 (redundant)','Add immediate-cash vs deferred-credit payment factor; auction-duration redundant as price is often pre-negotiated by middle agents; supports decentralised solar mini-processing.'],
 ['E04','Member, EDC Valley of Flowers, Chamoli','Weather risk & post-harvest shelf life','R_TIME_3 (negligible)','Add effect of frost/heavy rain on shelf life during loading delays; retail customer-service time has negligible supply-chain effect; subsidise standardised crates.'],
 ['E05','Member, EDC Valley of Flowers, Chamoli','Primary processing & grading at cluster','T_TIME_2 / T_QUAL_1 (overlap)','Add farm-cluster waxing/primary-processing capability; quality-inspection time overlaps initial quality assessment; lack of grading machines forces manual error.'],
 ['E06','Women SHG Leader, EDC Valley of Flowers, Chamoli','Women SHG labour in sorting/grading','M_COST_3 (redundant)','Add women-SHG labour charges in sorting/grading; mandi storage rent redundant as Malta rarely stays overnight; portable graders raise farm-gate returns.'],
 ['E07','Faculty, GBPUAT Hill Campus, Tehri Garhwal','Post-harvest physiology & transport loss','TR_TIME_5 (exclude from quality model)','Add physiological loss in weight (PLW) during long-haul; return-trip time should not enter fruit-quality models; a Garhwal-Malta GI tag would build brand equity.'],
 ['E08','Faculty, HNB Garhwal University, Srinagar','Agri-marketing & channel selection','R_COST_5 (does not drive channel choice)','Add digital price-awareness via mobile apps before selling; retail rent does not drive channel selection; evaluate e-commerce alongside mandi channels.'],
 ['E09','Operations Director, Devbhumi Agri FPO, Pauri Garhwal','FPO aggregation & working capital','F_TIME_2 + F_TIME_3 (merge)','Add access to institutional working-capital credit at peak harvest; combine the two post-harvest preparation-time factors; FPO can cut two middle tiers via NCR retail linkage.'],
 ['E10','Principal Scientist, ICAR-VPKAS (Almora/Chamoli unit)','Citrus crop protection & mandi quality','M_QUAL_4 (omit)','Add citrus greening / fruit-fly incidence on marketability; mandi grading-verification rarely rigorous, can be omitted; scientific orchard management plus reefer vans cut 30%+ losses.'],
 ['E11','Officer, Dept of Horticulture, Uttarkashi','Government subsidy & cluster logistics','TR_COST_5 (low relevance)','Add subsidy-disbursement efficiency for packaging; insurance/permits ignored by local pickup operators; NH-34/NH-07 cluster collection centres with cold rooms would transform logistics.'],
 ['E12','Representative, Garhwal Citrus FPO, Rudraprayag','Producer contracts & aggregation','R_TIME_2 (non-critical)','Add producer\u2013trader contractual purchase guarantee; retail display-arrangement time non-critical; panchayat-level aggregation prevents distress sales.'],
 ['E13','Faculty, UUHF Bharsar, Pauri Garhwal','Pomology & harvest maturity','F_COST_1 + F_COST_6 (merge)','Add fruit-maturity degree at harvest; merge the two farm labour-cost factors under total operational labour; premature harvest raises acidity rejection.'],
 ['E14','Official, APMC Mandi Board, Dehradun','Mandi operations & infrastructure','M_QUAL_1 (practically absent)','Add auction-bay availability at peak arrival; entry quality-inspection practically absent in current mandis; electronic weighing and cold yards optimise handling.'],
 ['E15','Representative, Alaknanda Organic FPO, Chamoli','Organic certification & premium realisation','T_COST_4 (duplicates initial grading)','Add organic-certification premium realisation; re-grading/re-packing cost duplicates initial grading cost; separate organic channels can yield ~40% higher margin.'],
 ['E16','Officer, Dept of Agriculture, Tehri Garhwal','Rural feeder-road access','TR_TIME_1 (moderate relevance)','Add feeder-road (PMGSY) availability to orchards; vehicle-availability time stable off-peak; feeder roads cut manual carrying and initial handling damage.'],
 ['E17','Scientist, Herbal Research Institute, Gopeshwar','By-product recovery & value addition','R_QUAL_5 (seldom measured)','Add by-product recovery value (peel oil, pulp); after-sales feedback quality seldom measured for fresh fruit; integrate processing with fresh-fruit channels.'],
 ['E18','Office-bearer, Uttarkashi Transport Union','Road transport operations','TR_COST_6 (merge with freight)','Add toll-gate delays and multi-axle entry restrictions; route/toll fees can merge with general freight; 3-tonne vehicles more efficient than large trucks on narrow routes.'],
 ['E19','Official, APMC Sub-Mandi, Rishikesh','Mandi turnaround & settlement','M_TIME_4 (uniform, APMC-governed)','Add vehicle turnaround time inside mandi; payment-settlement time uniform under APMC rules; direct farmer-to-retailer auction slots cut commission.'],
 ['E20','Faculty, Graphic Era University, Dehradun','Supply-chain modelling & MCDM','F_TIME_5 + F_TIME_6 (merge)','Add information asymmetry between hill farmers and terminal agents; merge loading and transit time into total outbound lead time; weight transport-risk highest via AHP/DEMATEL.'],
 ['E21','Federation Leader, Bhagirathi Mahila SHG, Uttarkashi','SHG micro-finance & packaging','R_COST_5 (external to rural chain)','Add micro-financing for reusable plastic crates; retail shop rent external to rural chain efficiency; corrugated boxes reduce downhill fruit damage.'],
 ['E22','Officer, Dept of Horticulture, Gopeshwar','Freight subsidy & harvest timing','T_TIME_4 (negligible)','Add freight-subsidy awareness; trader storage duration negligible (stock flips within 24h); harvest-timing awareness prevents low-brix fruit entering channels.'],
 ['E23','Operator, Chamoli Cold Chain & Logistics','Cold-chain & transport quality','TR_QUAL_2 (uniformly low at present)','Add reefer-van availability in remote districts; temperature-control capacity uniformly low in open trucks; phase in multi-temperature vehicles for mixed loads.'],
 ['E24','Scientist, Forest Research Institute (FRI), Dehradun','Sustainable packaging','R_TIME_3 (irrelevant for wholesale)','Add biodegradable-packaging adoption rate; retail customer-service time irrelevant for wholesale evaluation; pine-needle moulded boxes can replace plastic crates.'],
 ['E25','Official, Uttarakhand State Agri Marketing Board','Agri-marketing policy & e-NAM','M_COST_2 (group under entry fees)','Add inter-state border taxes and check-post clearance times; weighment charges minimal, group under entry fees; accelerate e-NAM integration across hill mandis.'],
 ['E26','Scientist, KVK Pokhari, Chamoli','Extension & post-harvest handling','F_COST_5 (irrelevant \u2014 no farm storage)','Add farmer training in post-harvest handling and washing; on-farm storage cost irrelevant as hill farmers lack storage; field sorting cuts freight on unmarketable fruit.'],
 ['E27','Office-bearer, Yamuna Valley Farmers Co-op','Cooperative marketing & bargaining','T_COST_5 (overlaps mandi charges)','Add bargaining power of co-ops vs regional traders; commission/agency fee overlaps mandi charges; direct institutional procurement is most profitable channel.'],
 ['E28','Faculty, Graphic Era Hill University, Dehradun','Consumer behaviour & branding','R_TIME_2 (statistically insignificant)','Add consumer willingness-to-pay for authentic hill Malta; display-layout time statistically insignificant; "Himalayan pristine" branding can unlock premium pricing.'],
 ['E29','Representative, Garhwal Logistics Association','Freight economics','TR_QUAL_4 (combine with vehicle condition)','Add fuel-price volatility surcharge in winter; transport-safety measures can combine with vehicle condition; aggregating freight across 3\u20134 villages cuts per-quintal cost.'],
 ['E30','Faculty, UUHF Ranichauri, Tehri Garhwal','Pomology & rootstock','M_QUAL_5 (redundant with transport loading)','Add rootstock influence on shelf life and skin thickness; post-auction handling redundant with transport-loading variables; disease-free rootstock simplifies grading.'],
 ['E31','Officer, State Agriculture Dept, Pauri Garhwal','Crop/post-harvest risk & insurance','F_COST_5 (rarely stored on farm)','Add post-harvest-loss insurance coverage; on-farm storage cost minor as farmers rarely store; extend crop insurance to weather-related transit losses.'],
 ['E32','Representative, Mandakini Valley Producer Co.','Producer logistics & aggregation hubs','T_TIME_3 (duplicates farm sorting)','Add reefer-hub availability at highway entry (Rishikesh); trader re-sorting time duplicates farm-level sorting; hub-and-spoke in river valleys streamlines logistics.'],
 ['E33','Faculty, Doon University, Dehradun','Market economics & demand','R_COST_5 (adds noise)','Add demand elasticity across tier-1 wholesale markets; retail space rent adds noise to cost analysis; optimise direct logistics to Delhi-NCR (Azadpur).'],
 ['E34','Operator, Hill Route Carriers, Srinagar','Transport operations & safety','TR_TIME_4 (combine with destination transit)','Add driver fatigue and night-driving restrictions; unloading time static, combine with destination transit; staggered departures prevent valley bottlenecks.'],
 ['E35','Scientist, KVK Chinyalisaur, Uttarkashi','Post-harvest technology','M_TIME_2 (minor in small mandis)','Add ethylene absorbers / anti-microbial coatings; auction-waiting time minor in small hill mandis; wash-and-wax extends Malta shelf life from ~15 to ~45 days.'],
 ['E36','Official, APMC Mandi, Kotdwar','Mandi governance & auctions','M_COST_2 (group under admin overhead)','Add transparency of open-cry vs secret bidding; weighbridge fee nominal, group under admin overhead; digital weighbridges linked to e-NAM ensure trusted weights.'],
 ['E37','Faculty, HNB Garhwal University','Labour economics & gender','R_TIME_4 (hard to capture in informal retail)','Add gender parity in harvesting/sorting wages; retail stock-turnover hard to capture in informal retail; standardise wages and train rural youth in logistics.'],
 ['E38','Office-bearer, Kedarnath Agri Producers Co.','Inter-state logistics & documentation','T_COST_3 (overlaps retail cost)','Add customs/transit documentation for inter-state vehicles; agent storage-and-handling fee overlaps retail cost; pre-booked co-op logistics cut shipping ~18%.'],
 ['E39','Officer, Dept of Horticulture, Tehri Garhwal','Storage infrastructure & energy','F_TIME_4 (merge into transport arrival)','Add power supply for local cold stores; farm waiting-for-transport time can merge into transport-arrival time; grid instability makes solar cold stores essential.'],
 ['E40','Office-bearer, Joshimath Goods Carrier Union','High-altitude transport support','TR_COST_5 (static annual, not per-trip)','Add vehicle repair/breakdown support along high-altitude routes; insurance/permits are static annual, not per-trip; breakdown/refuel points on Badrinath NH prevent losses.'],
 ['E41','Faculty, Graphic Era University, Dehradun','SCM risk & quantitative methods','R_QUAL_4 (subjective; prefer hard metrics)','Add supply-chain risk-mitigation (weather/political/market); customer quality-perception subjective; apply SEM to extract latent supply-chain barriers.'],
 ['E42','Representative, Himalayan Organic Farmer Co.','Fair-trade & producer welfare','T_COST_4 (eliminate via field packing)','Add fair-trade pricing guarantees for small/marginal farmers; repacking cost should be eliminated by standardising field packing; village collection centres bypass middle traders.'],
 ['E43','Official, APMC Sub-Mandi, Haldwani/Dehradun','Mandi logistics & cross-docking','M_QUAL_2 (long storage rare)','Add cross-docking at gateway terminals (Rishikesh/Haridwar); mandi storage-condition minor as produce moves fast; cross-docking consolidates hill pickups into long-haul trucks.'],
 ['E44','Faculty, UUHF Ranichauri, Tehri Garhwal','Grading science & post-harvest decay','R_COST_6 (merge with shop operating cost)','Add effect of grading accuracy on storage decay; retail handling labour can merge with shop operating cost; cluster-level sizing machines preserve batch quality.'],
 ['E45','Operator, Garhwal Agro Express, Rishikesh','Transport technology & routing','TR_TIME_5 (irrelevant to forward delivery)','Add real-time GPS tracking / route-optimisation adoption; return-trip time irrelevant to forward delivery efficiency; tracking cuts unscheduled halts on the 200 km hill corridor.'],
]
para('Table 4.10a  Expert validation panel: designation/affiliation, area of expertise, factors validated or flagged, and specific '
     'feedback (N = 45).', italic=True, size=9)
make_table(
    ['ID','Designation / Affiliation','Area of expertise','Factor(s) validated / flagged','Specific suggestion / feedback'],
    experts,
    widths=[0.4, 1.8, 1.5, 1.35, 3.05], font_size=7)
source('Source: Author\u2019s expert-validation exercise, Garhwal region. Expert identities are recorded in the researcher\u2019s master '
       'register; codes E01\u2013E45 are used here in place of names for the thesis annexure. Factor codes map to the five-stage '
       '(F, M, T, TR, R) \u00d7 three-dimension (COST, TIME, QUAL) scheme of Section 4.11.')

H('Note on the two respondent lists', level=2)
para('Your records contain two related but distinct datasets. The coded register MGR1236\u2013MGR1300 (65 entries) is the '
     '"Experts/Officials" stakeholder stratum within the main survey of N = 1,300 (Table 4.5). The 45 named domain specialists above '
     'form the content-validity panel whose qualitative Q1\u2013Q3 feedback refined the factor pool. Keep the two clearly separated in '
     'the thesis: the 65-row register belongs with the sampling tables in Chapter 4/5, and the 45-expert validation table belongs in '
     'Section 4.7.3 / Annexure B. [CONFIRM whether any individuals appear in both, and reconcile the counts in your response to the DRC.]',
     align='justify')

# ============================================================= SECTION 4
doc.add_page_break()
H('Section 4 — Justification of Block Identification and Sample-Size Adequacy', level=1)
para('Suggested placement: expand Section 4.8 (Sampling Design) with a new sub-section 4.8.1 "Identification of study blocks", and '
     'strengthen Section 4.9 (Sample Size Justification). Placeholders marked [INSERT]/[CONFIRM] require your own block-selection '
     'records and figures.', italic=True, size=10)

H('4.8.1 Identification of the study blocks', level=2)
para('The four study districts \u2014 Pauri Garhwal, Tehri Garhwal, Rudraprayag and Chamoli \u2014 were selected because they '
     'constitute the principal Malta-growing belt of Garhwal and together capture the region\u2019s full range of terrain, road '
     'connectivity and market access (Directorate of Horticulture and Food Processing, Uttarakhand, 2023; Kumar and Bisht, 2020). '
     'Within each district, development blocks were not chosen arbitrarily; they were identified through a transparent, '
     'multi-criteria procedure so that the surveyed blocks would be genuinely representative of Malta production and its distribution '
     'conditions rather than merely accessible.', align='justify')
para('Blocks were shortlisted against five criteria, applied in sequence:', align='justify')
numbered('Malta production concentration. Blocks with the highest recorded area and output under Malta were prioritised, using '
         'District Horticulture Office records and the Uttarakhand at a Glance (2021\u201322) compendium as the frame. '
         '[INSERT the block-wise area/production figures you used.]')
numbered('Agro-climatic / altitude representativeness. Blocks were chosen to span the sub-tropical valley to mid-temperate gradient '
         'along which Malta is grown, so that low-, mid- and higher-altitude orchard conditions are all represented.')
numbered('Road-connectivity and market-access gradient. Blocks were selected to include well-connected, moderately connected and '
         'remote/roadhead-deficient conditions, because connectivity is a primary driver of cost, time and quality outcomes in the channel.')
numbered('Presence of channel actors and institutions. Blocks hosting the full set of stakeholders \u2014 growers, village traders, '
         'commission agents, mandis/sub-mandis, transporters, retailers, FPOs, cooperatives, EDCs and women SHGs \u2014 were preferred '
         'so that every stage of the channel could be observed within the same block context.')
numbered('Logistical feasibility and safety. Given the fragile mountain terrain and monsoon/landslide risk, blocks were finalised '
         'subject to safe seasonal accessibility, documented through geo-tagged field visits.')
para('Applying these criteria produced the following set of surveyed blocks, each confirmed by geo-tagged field evidence '
     '(Annexure C). [CONFIRM the block names and add any omitted blocks; the entries below are inferred from your geo-tagged photographs.]',
     align='justify')
make_table(
    ['District','Identified blocks (surveyed)','Illustrative field sites (geo-tagged)','Basis for selection'],
    [
     ['Pauri Garhwal','Thailisain; Chaubatta Khal [CONFIRM others]','Kainyur/Raull; Jantadevi','Established growing pockets; better connectivity; active FPO/SHG presence'],
     ['Tehri Garhwal','Chamba; Jakhnidhar; Pratapnagar [CONFIRM]','Sabli/Ranichauri; Tharsalgaon; Petouri','Significant production; steep dispersed orchards; connectivity gradient'],
     ['Rudraprayag','[INSERT blocks, e.g., Jakholi/Ukhimath]','[INSERT field sites]','Emerging production; high-altitude, road-constrained conditions'],
     ['Chamoli','Pokhari; Joshimath [CONFIRM others]','Sankri/Nail; Rupung/Molta','Substantial production; remote, long market distances; transport risk'],
    ],
    widths=[1.2, 1.9, 1.9, 2.1], font_size=8)
source('Source: Author, from District Horticulture Office records, Uttarakhand at a Glance (2021\u201322), and geo-tagged field '
       'enumeration (Annexure C).')

H('4.9 (strengthened) Adequacy and justification of the sample size (N = 1,300)', level=2)
para('The study rests on 1,300 valid responses. Its adequacy is defended on five grounds, three of which restate the existing '
     'argument and two of which respond directly to the committee\u2019s query.', align='justify')
numbered('Stratified-design requirement. The design is multi-stakeholder and multi-district: five operational stakeholder groups '
         'across four districts generate 20 district-by-stakeholder cells (plus the expert stratum). Cochran\u2019s (1977) logic for '
         'stratified sampling requires the total sample to grow with the number of strata if within-cell estimates are to stay stable; '
         '1,300 satisfies this comfortably, with no analytic cell reduced to a fragile count (see Table 4.5).')
numbered('Statistical robustness. Reliability coefficients, inter-item correlations and the descriptive statistics that feed the '
         'decision matrices are estimated far more precisely at N = 1,300 than in a small panel; sampling error is correspondingly '
         'small, and the fuzzy weights and outranking relations inherit this stability (Hair et al., 2019; Tabachnick and Fidell, 2019).')
numbered('External validity. The Garhwal Malta channel is heterogeneous, spanning smallholders, intermediaries and logistics '
         'providers under very different constraints; a large stratified sample lets the findings generalise across the belt rather than '
         'to a narrow subset (Bryman, 2016; Saunders et al., 2019).')
numbered('Sufficiency against formal thresholds. For a finite but large grower-and-actor population, Krejcie and Morgan\u2019s (1970) '
         'and Cochran\u2019s (1977) formulae return a required sample of roughly 380\u2013400 at 95% confidence and \u00b15% precision. '
         'At N = 1,300 the study exceeds this several times over, so precision is well within the conventional margin even after '
         'district-and-stakeholder stratification. [INSERT your exact population frame estimate and the computed n* if available.]')
numbered('Overcoming the expert-only limitation of MCDM. Fuzzy AHP/ELECTRE applications often rest on a handful of experts, which makes '
         'the weights sensitive to a few individuals. By grounding perception measurement in a large multi-actor sample and reserving '
         'expert judgement for the pairwise-comparison stage, the study reduces dependence on any small panel (Saaty, 1980; '
         'Ishizaka and Nemery, 2013).')
para('Transparency on limits. No complete official enumeration of Malta households and traders exists for the region, so the sampling '
     'frame is an assembled approximation; formal response/non-response rates were not logged; and the expert sub-sample was purposive. '
     'These caveats are recorded in Section 4.9.1 and Section 8.5, and claims of generalisability are made accordingly.', align='justify')

# ============================================================= SECTION 5
doc.add_page_break()
H('Section 5 — DRC Submission: Point-by-Point Compliance Statement', level=1)
para('This is the cover response for formal DRC submission. Replace the "Committee observation" wording with the exact text from your '
     'DRC minutes, and confirm every cross-reference before submitting.', italic=True, size=10)

para('To: The Chairperson and Members, Doctoral Research Committee', bold=True, size=11, space_after=2)
para('From: Mohit Gundwal, Research Scholar (Management)', size=11, space_after=2)
para('Subject: Compliance report on revisions carried out in response to DRC observations', size=11, space_after=2)
para('Date: [INSERT] (within the one-month revision window)', size=11, space_after=10)

para('Respected Committee Members, the observations recorded in the DRC meeting have been addressed in full. Each observation, the '
     'action taken, and the exact location of the revision in the thesis are set out below. A revised copy of the thesis with the '
     'changes incorporated accompanies this report.', align='justify')

make_table(
    ['#','Committee observation (as minuted)','Action taken / revision made','Location in revised thesis'],
    [
     ['1','Conclusions should be made region-specific to Garhwal rather than generic.',
      'Added a dedicated set of six region-specific conclusions plus a district-level texture table (Pauri, Tehri, Rudraprayag, Chamoli), tying the cost-dominance, upstream-criticality and transportation-model results to the Garhwal belt.',
      'Chapter 8, new Section 8.1a; Table 8.1a'],
     ['2','Geo-tagged field photographs must be presented in proper academic format.',
      'Introduced Plate-series conventions (numbering, captions with decimal-degree coordinates, date/time, datum and source), added a "List of Plates", a geo-tag register (Table C.1), and placed the images in Annexure C; representative plates cross-referenced in Chapters 4\u20135.',
      'Section 4.6.1 (note); Annexure C; front-matter List of Plates'],
     ['3','The expert panel and their feedback must be documented as a formal validation.',
      'Compiled the 45-member content-validity panel into a structured table (ID, designation/affiliation, area of expertise, factors validated/flagged, specific feedback), and linked it to the 144\u219282 factor-refinement process.',
      'Section 4.7.3; Annexure B; cross-linked to Section 4.11'],
     ['4','Justify how the study blocks were identified.',
      'Added Section 4.8.1 setting out the five-criterion block-identification procedure (production concentration, agro-climatic representativeness, connectivity gradient, presence of channel actors, feasibility/safety), with a district-wise block table validated by geo-tagged evidence.',
      'Section 4.8.1; Table 4.8.1'],
     ['5','Justify the adequacy of the data / sample size.',
      'Strengthened Section 4.9 with five grounds, including an explicit Krejcie\u2013Morgan/Cochran threshold check showing N = 1,300 exceeds the required ~380\u2013400, plus transparent treatment of frame, non-response and purposive-expert limitations.',
      'Section 4.9; Section 4.9.1; Section 8.5'],
     ['6','[INSERT any further observation from the minutes]',
      '[INSERT action taken]',
      '[INSERT location]'],
    ],
    widths=[0.35, 2.15, 3.1, 1.5], font_size=8)
source('Source: Author, mapped to the DRC meeting minutes dated [INSERT].')

para('I remain available to present any of these revisions to the Committee and to make further changes as directed. I am grateful '
     'for the Committee\u2019s guidance, which has materially improved the rigour and regional specificity of the thesis.', align='justify')
para('[Signature]', space_after=2)
para('Mohit Gundwal', bold=True, space_after=2)
para('Research Scholar (Management)   |   Supervisor: [INSERT]   |   Co-Supervisor: [INSERT]', size=10)

# ---- footer note on figures/numbers ----
doc.add_page_break()
H('Checklist before you submit', level=1)
for item in [
 'Confirm all coordinates, dates and block names against your GPS metadata and field diary.',
 'Reconcile the 65-row expert/official register (MGR1236\u2013MGR1300) with the 45-member validation panel; explain the relationship in the DRC report.',
 'Replace every [INSERT]/[CONFIRM] placeholder with your verified figure or name.',
 'Insert the exact wording of each DRC observation from the official minutes into Section 5.',
 'Re-read Section 1 and 4 in your own voice so they match the register of your existing chapters.',
 'Verify the two possibly-duplicated landslide-road photographs before including both.',
 'Add the Krejcie\u2013Morgan computed n* and your population-frame estimate if your supervisor wants the arithmetic shown.',
]:
    bullet(item)

doc.save('DRC_Revision_Package.docx')
print('SAVED DRC_Revision_Package.docx')
print('paragraphs:', len(doc.paragraphs), 'tables:', len(doc.tables))
