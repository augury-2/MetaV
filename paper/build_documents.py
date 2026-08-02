"""
Build downloadable document versions of the manuscript and appendices.

Produces, in paper/dist/:
  * <name>.html  - self-contained single file with all figures embedded as base64.
                   Opens in any browser; use the browser's Print > Save as PDF for a PDF.
  * <name>.docx  - Word document with figures embedded, for editing and submission.

No external tooling required (no pandoc, no LaTeX). Handles the Markdown subset actually
used in these documents: ATX headings, paragraphs with inline bold/italic/code, pipe tables,
blockquotes, fenced code blocks, ordered and unordered lists, horizontal rules, and the
`**[Figure N: `path`]**` figure markers used in the manuscript.

Usage:
    python build_documents.py
"""
from __future__ import annotations

import base64
import html
import os
import re
import sys

DOC_DIR = os.path.dirname(os.path.abspath(__file__))
DIST = os.path.join(DOC_DIR, "dist")
os.makedirs(DIST, exist_ok=True)

DOCS = [
    ("MANUSCRIPT.md", "Marketing Without Growth"),
    ("WEB_APPENDIX.md", "Web Appendix"),
    ("POSITIONING_MEMO.md", "Positioning and Reviewer-Risk Memo"),
]

FIG_MARKER = re.compile(r"^\*\*\[Figure\s*([^:]*):\s*`([^`]+)`\]\*\*\s*$")


# ======================================================================================
# Block-level parsing shared by both writers
# ======================================================================================
def parse_blocks(md: str, base_dir: str):
    """Yield (kind, payload) blocks. kind in: h, p, table, quote, code, ul, ol, hr, figure."""
    lines = md.split("\n")
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        m = FIG_MARKER.match(stripped)
        if m:
            path = os.path.normpath(os.path.join(base_dir, m.group(2)))
            yield ("figure", (m.group(1).strip(), path))
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            yield ("code", "\n".join(buf))
            continue

        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", stripped):
            yield ("hr", None)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            yield ("h", (len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        # pipe table: current line and next line is a separator row
        if stripped.startswith("|") and i + 1 < n and re.match(
                r"^\|[\s:|-]+\|?\s*$", lines[i + 1].strip()):
            rows = []
            header = _split_row(stripped)
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i].strip()))
                i += 1
            yield ("table", (header, rows))
            continue

        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            yield ("quote", "\n".join(buf))
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            items = []
            while i < n:
                mm = re.match(r"^\s*\d+\.\s+(.*)$", lines[i])
                if not mm:
                    if lines[i].strip() and lines[i].startswith(("   ", "\t")) and items:
                        items[-1] += " " + lines[i].strip()
                        i += 1
                        continue
                    break
                items.append(mm.group(1).strip())
                i += 1
            yield ("ol", items)
            continue

        if re.match(r"^[-*+]\s+", stripped):
            items = []
            while i < n:
                mm = re.match(r"^\s*[-*+]\s+(.*)$", lines[i])
                if not mm:
                    if lines[i].strip() and lines[i].startswith(("   ", "\t")) and items:
                        items[-1] += " " + lines[i].strip()
                        i += 1
                        continue
                    break
                items.append(mm.group(1).strip())
                i += 1
            yield ("ul", items)
            continue

        # paragraph: accumulate until blank line or a new block starts
        buf = []
        while i < n and lines[i].strip():
            s = lines[i].strip()
            if (re.match(r"^#{1,6}\s", s) or s.startswith("```") or s.startswith("|")
                    or s.startswith(">") or re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", s)
                    or FIG_MARKER.match(s)):
                break
            buf.append(s)
            i += 1
        if buf:
            yield ("p", " ".join(buf))


def _split_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


# ======================================================================================
# HTML writer
# ======================================================================================
CSS = """
:root { --ink:#1a1a1a; --muted:#5a5a5a; --rule:#d8d8d8; --accent:#1b6ca8; --warn:#c0392b; }
* { box-sizing: border-box; }
body { font-family: Georgia, 'Iowan Old Style', 'Times New Roman', serif; color: var(--ink);
       max-width: 46rem; margin: 0 auto; padding: 3rem 1.5rem 6rem; line-height: 1.62;
       font-size: 16px; }
h1 { font-size: 1.85rem; line-height: 1.25; margin: 2.5rem 0 1rem; }
h1:first-child { margin-top: 0; }
h2 { font-size: 1.35rem; margin: 2.6rem 0 .8rem; padding-bottom: .3rem;
     border-bottom: 2px solid var(--rule); }
h3 { font-size: 1.12rem; margin: 2rem 0 .6rem; }
h4 { font-size: 1rem; margin: 1.5rem 0 .5rem; color: var(--muted); }
p { margin: .85rem 0; }
a { color: var(--accent); }
hr { border: none; border-top: 1px solid var(--rule); margin: 2.5rem 0; }
blockquote { margin: 1.4rem 0; padding: .9rem 1.2rem; background: #f6f8fa;
             border-left: 4px solid var(--accent); font-size: .95rem; }
blockquote p:first-child { margin-top: 0; } blockquote p:last-child { margin-bottom: 0; }
code { font-family: ui-monospace, 'SF Mono', Menlo, Consolas, monospace; font-size: .88em;
       background: #f0f2f4; padding: .1em .35em; border-radius: 3px; }
pre { background: #f6f8fa; padding: .9rem 1.1rem; overflow-x: auto; border-radius: 4px;
      border: 1px solid var(--rule); }
pre code { background: none; padding: 0; font-size: .82rem; line-height: 1.45; }
table { border-collapse: collapse; width: 100%; margin: 1.5rem 0; font-size: .86rem;
        font-family: -apple-system, 'Segoe UI', Helvetica, Arial, sans-serif; }
th, td { border: 1px solid var(--rule); padding: .45rem .6rem; text-align: left;
         vertical-align: top; }
th { background: #f0f2f4; font-weight: 600; }
tr:nth-child(even) td { background: #fafbfc; }
figure { margin: 2rem 0; text-align: center; page-break-inside: avoid; }
figure img { max-width: 100%; height: auto; border: 1px solid var(--rule); }
figcaption { font-size: .82rem; color: var(--muted); margin-top: .5rem;
             font-family: -apple-system, 'Segoe UI', Helvetica, Arial, sans-serif; }
ul, ol { margin: .85rem 0; padding-left: 1.6rem; }
li { margin: .3rem 0; }
.banner { background:#fff8e1; border:1px solid #f0d48a; border-left:5px solid #d9a406;
          padding: 1rem 1.2rem; margin: 2rem 0; font-size:.93rem; }
.toc { background:#f6f8fa; border:1px solid var(--rule); padding:1rem 1.4rem; margin:2rem 0;
       font-family:-apple-system,'Segoe UI',Helvetica,Arial,sans-serif; font-size:.88rem; }
.toc ol { padding-left: 1.3rem; } .toc a { text-decoration: none; }
@media print {
  body { max-width: none; padding: 0; font-size: 10.5pt; }
  h1,h2,h3 { page-break-after: avoid; }
  table, figure, blockquote, pre { page-break-inside: avoid; }
  a { color: var(--ink); text-decoration: none; }
  .noprint { display: none; }
}
"""


# --------------------------------------------------------------------------------------
# Inline parsing (shared by the HTML and DOCX writers)
# --------------------------------------------------------------------------------------
# A recursive tokenizer is required rather than a chain of regex substitutions, because this
# manuscript uses (a) backslash-escaped asterisks inside emphasis, as in `*S\**` for the
# italic flip-point symbol, and (b) bold spans that contain italics or code, as in
# `***p* = .751**`. Flat substitution mishandles both.

Span = tuple[str, frozenset]          # (text, styles) where styles ⊆ {bold, italic, code, del}


def parse_inline(text: str, styles: frozenset = frozenset()) -> list[Span]:
    """Tokenize inline markdown into styled spans, recursing into nested emphasis."""
    spans: list[Span] = []
    i, n = 0, len(text)
    buf: list[str] = []

    def flush():
        if buf:
            spans.append(("".join(buf), styles))
            buf.clear()

    while i < n:
        ch = text[i]

        # backslash escape: the next character is literal
        if ch == "\\" and i + 1 < n:
            buf.append(text[i + 1])
            i += 2
            continue

        # code span (never contains markup)
        if ch == "`":
            j = text.find("`", i + 1)
            if j != -1:
                flush()
                spans.append((text[i + 1:j], styles | {"code"}))
                i = j + 1
                continue

        # strong: **...**
        if text.startswith("**", i):
            j = _find_close(text, i + 2, "**")
            if j != -1:
                flush()
                spans.extend(parse_inline(text[i + 2:j], styles | {"bold"}))
                i = j + 2
                continue

        # strikethrough
        if text.startswith("~~", i):
            j = _find_close(text, i + 2, "~~")
            if j != -1:
                flush()
                spans.extend(parse_inline(text[i + 2:j], styles | {"del"}))
                i = j + 2
                continue

        # emphasis: *...*
        if ch == "*":
            j = _find_close(text, i + 1, "*")
            if j != -1 and j > i + 1:
                flush()
                spans.extend(parse_inline(text[i + 1:j], styles | {"italic"}))
                i = j + 1
                continue

        buf.append(ch)
        i += 1

    flush()
    return spans


def _find_close(text: str, start: int, delim: str) -> int:
    """Index of the next unescaped `delim`, skipping code spans. -1 if absent."""
    i, n, d = start, len(text), len(delim)
    while i < n:
        if text[i] == "\\":
            i += 2
            continue
        if text[i] == "`":
            j = text.find("`", i + 1)
            i = (j + 1) if j != -1 else n
            continue
        if text.startswith(delim, i):
            # for single-char '*', do not treat the '**' of a strong marker as a close
            if d == 1 and text.startswith("**", i):
                i += 2
                continue
            # Disambiguate an odd run of asterisks, e.g. `*S\**` immediately followed by a
            # closing `**`, which yields `***`. The inner emphasis must claim the first
            # asterisk, so the strong delimiter starts one character later.
            if d == 2:
                run = len(text[i:]) - len(text[i:].lstrip("*"))
                if run == 3:
                    return i + 1
            return i
        i += 1
    return -1


def _spans_to_html(s: str) -> str:
    out = []
    for text, styles in parse_inline(s):
        esc = html.escape(text, quote=False)
        if "code" in styles:
            esc = f"<code>{esc}</code>"
        if "italic" in styles:
            esc = f"<em>{esc}</em>"
        if "bold" in styles:
            esc = f"<strong>{esc}</strong>"
        if "del" in styles:
            esc = f"<del>{esc}</del>"
        out.append(esc)
    return "".join(out)


def inline_html(s: str) -> str:
    """Inline markdown -> HTML, with markdown links rendered as anchors."""
    parts, pos = [], 0
    for m in re.finditer(r"(?<!\\)\[([^\]]+)\]\(([^)]+)\)", s):
        parts.append(_spans_to_html(s[pos:m.start()]))
        parts.append(f'<a href="{html.escape(m.group(2), quote=True)}">'
                     f'{_spans_to_html(m.group(1))}</a>')
        pos = m.end()
    parts.append(_spans_to_html(s[pos:]))
    return "".join(parts)


def embed_image(path: str) -> str | None:
    if not os.path.exists(path):
        return None
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("ascii")


def slug(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", re.sub(r"[*`\[\]]", "", text).lower()).strip("-")


def to_html(md: str, title: str, base_dir: str) -> str:
    out = [f"<!DOCTYPE html><html lang='en'><head><meta charset='utf-8'>",
           f"<meta name='viewport' content='width=device-width,initial-scale=1'>",
           f"<title>{html.escape(title)}</title><style>{CSS}</style></head><body>"]
    headings: list[tuple[int, str, str]] = []
    body: list[str] = []
    missing: list[str] = []

    for kind, payload in parse_blocks(md, base_dir):
        if kind == "h":
            lvl, text = payload
            sid = slug(text)
            if lvl == 2:
                headings.append((lvl, text, sid))
            body.append(f"<h{lvl} id='{sid}'>{inline_html(text)}</h{lvl}>")
        elif kind == "p":
            body.append(f"<p>{inline_html(payload)}</p>")
        elif kind == "hr":
            body.append("<hr>")
        elif kind == "code":
            body.append(f"<pre><code>{html.escape(payload)}</code></pre>")
        elif kind == "quote":
            inner = "".join(
                f"<p>{inline_html(par)}</p>" for par in re.split(r"\n\s*\n", payload)
                if par.strip())
            cls = " class='banner'" if "Data status" in payload else ""
            body.append(f"<blockquote{cls}>{inner}</blockquote>")
        elif kind == "ul":
            body.append("<ul>" + "".join(f"<li>{inline_html(x)}</li>" for x in payload) + "</ul>")
        elif kind == "ol":
            body.append("<ol>" + "".join(f"<li>{inline_html(x)}</li>" for x in payload) + "</ol>")
        elif kind == "table":
            header, rows = payload
            th = "".join(f"<th>{inline_html(c)}</th>" for c in header)
            trs = "".join(
                "<tr>" + "".join(f"<td>{inline_html(c)}</td>" for c in r) + "</tr>"
                for r in rows)
            body.append(f"<table><thead><tr>{th}</tr></thead><tbody>{trs}</tbody></table>")
        elif kind == "figure":
            label, path = payload
            b64 = embed_image(path)
            if b64:
                body.append(
                    f"<figure><img alt='Figure {html.escape(label)}' "
                    f"src='data:image/png;base64,{b64}'>"
                    f"<figcaption>Figure {html.escape(label)}</figcaption></figure>")
            else:
                missing.append(path)
                body.append(f"<p><em>[Figure {html.escape(label)} not found: "
                            f"{html.escape(path)}]</em></p>")

    if len(headings) > 3:
        toc = "".join(f"<li><a href='#{sid}'>{inline_html(t)}</a></li>" for _, t, sid in headings)
        out.append(f"<nav class='toc noprint'><strong>Contents</strong><ol>{toc}</ol></nav>")
    out.extend(body)
    out.append("</body></html>")
    if missing:
        print(f"    warning: {len(missing)} figure(s) not found")
    return "\n".join(out)


# ======================================================================================
# DOCX writer
# ======================================================================================
def to_docx(md: str, title: str, base_dir: str, out_path: str):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.25

    def add_runs(par, text: str):
        """Render inline markdown into runs on an existing paragraph."""
        # flatten markdown links to "label (url)" so the text stays readable in Word
        text = re.sub(r"(?<!\\)\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
        for content, styles in parse_inline(text):
            if not content:
                continue
            r = par.add_run(content)
            r.bold = "bold" in styles
            r.italic = "italic" in styles
            if "del" in styles:
                r.font.strike = True
            if "code" in styles:
                r.font.name = "Consolas"
                r.font.size = Pt(9.5)

    missing = []
    for kind, payload in parse_blocks(md, base_dir):
        if kind == "h":
            lvl, text = payload
            p = doc.add_heading(level=min(lvl, 4))
            add_runs(p, text)
        elif kind == "p":
            add_runs(doc.add_paragraph(), payload)
        elif kind == "hr":
            p = doc.add_paragraph()
            p.add_run("_" * 60).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "code":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif kind == "quote":
            for par in re.split(r"\n\s*\n", payload):
                if par.strip():
                    p = doc.add_paragraph(style="Intense Quote")
                    add_runs(p, par.strip().replace("\n", " "))
        elif kind in ("ul", "ol"):
            style = "List Bullet" if kind == "ul" else "List Number"
            for item in payload:
                add_runs(doc.add_paragraph(style=style), item)
        elif kind == "table":
            header, rows = payload
            ncol = len(header)
            tbl = doc.add_table(rows=1, cols=ncol)
            tbl.style = "Light Grid Accent 1"
            for j, c in enumerate(header):
                cell = tbl.rows[0].cells[j]
                cell.text = ""
                add_runs(cell.paragraphs[0], c)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = tbl.add_row().cells
                for j in range(ncol):
                    cells[j].text = ""
                    add_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "")
            for r in tbl.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(8.5)
                            run.font.name = "Calibri"
        elif kind == "figure":
            label, path = payload
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(6.1))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(f"Figure {label}")
                r.italic = True
                r.font.size = Pt(9)
            else:
                missing.append(path)
    doc.save(out_path)
    if missing:
        print(f"    warning: {len(missing)} figure(s) not found")


# ======================================================================================
def main():
    print("Building downloadable documents ...")
    for fname, title in DOCS:
        src = os.path.join(DOC_DIR, fname)
        if not os.path.exists(src):
            print(f"  skip {fname} (not found)")
            continue
        md = open(src, encoding="utf-8").read()
        stem = os.path.splitext(fname)[0]

        html_path = os.path.join(DIST, f"{stem}.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(to_html(md, title, DOC_DIR))
        print(f"  {os.path.relpath(html_path, DOC_DIR)}  "
              f"({os.path.getsize(html_path) / 1e6:.2f} MB, figures embedded)")

        docx_path = os.path.join(DIST, f"{stem}.docx")
        to_docx(md, title, DOC_DIR, docx_path)
        print(f"  {os.path.relpath(docx_path, DOC_DIR)}  "
              f"({os.path.getsize(docx_path) / 1e6:.2f} MB)")

    print("\nTo produce a PDF: open the .html file in a browser and use Print > Save as PDF.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
