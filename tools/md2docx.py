#!/usr/bin/env python3
"""Minimal Markdown -> DOCX converter for manuscript deliverables.

Supports: ATX headings, paragraphs, bold/italic/code inline spans, pipe tables,
blockquotes, unordered and ordered lists, horizontal rules.

Usage: python3 md2docx.py input.md output.docx "Document Title"
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INLINE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`.+?`|\[.+?\]\(.+?\))")
LINK = re.compile(r"^\[(.+?)\]\((.+?)\)$")


def add_runs(paragraph, text):
    """Add inline-formatted runs to a paragraph."""
    text = text.replace("\\*", "\u0001")
    for piece in INLINE.split(text):
        if not piece:
            continue
        piece = piece.replace("\u0001", "*")
        bold = italic = mono = False
        link = LINK.match(piece)
        if link:
            piece = f"{link.group(1)} ({link.group(2)})"
        elif piece.startswith("***") and piece.endswith("***"):
            piece, bold, italic = piece[3:-3], True, True
        elif piece.startswith("**") and piece.endswith("**"):
            piece, bold = piece[2:-2], True
        elif piece.startswith("*") and piece.endswith("*"):
            piece, italic = piece[1:-1], True
        elif piece.startswith("`") and piece.endswith("`"):
            piece, mono = piece[1:-1], True
        run = paragraph.add_run(piece)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)


def flush_table(doc, rows):
    """Render collected pipe-table rows as a DOCX table."""
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [r for r in cells if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in r)]
    if not cells:
        return
    width = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=width)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j in range(width):
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, row[j] if j < len(row) else "")
            for run in para.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
    doc.add_paragraph()


def convert(md_path, docx_path, title):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    lines = open(md_path, encoding="utf-8").read().splitlines()
    table_buf, para_buf, first_heading = [], [], True

    def flush_para():
        if para_buf:
            add_runs(doc.add_paragraph(), " ".join(para_buf))
            para_buf.clear()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|"):
            flush_para()
            table_buf.append(stripped)
            i += 1
            continue
        if table_buf:
            flush_table(doc, table_buf)
            table_buf = []

        if not stripped:
            flush_para()
        elif re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", stripped):
            flush_para()
        elif stripped.startswith("#"):
            flush_para()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if first_heading:
                doc.add_paragraph()
                first_heading = False
            head = doc.add_heading(level=min(level, 4))
            add_runs(head, text)
            for run in head.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            if level == 1:
                head.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith(">"):
            flush_para()
            para = doc.add_paragraph(style="Intense Quote")
            add_runs(para, stripped.lstrip("> ").strip())
        elif re.match(r"^[-*+]\s+", stripped):
            flush_para()
            para = doc.add_paragraph(style="List Bullet")
            add_runs(para, re.sub(r"^[-*+]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            flush_para()
            para = doc.add_paragraph(style="List Number")
            add_runs(para, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            para_buf.append(stripped)
        i += 1

    flush_para()
    if table_buf:
        flush_table(doc, table_buf)

    doc.core_properties.title = title
    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else "Document")
